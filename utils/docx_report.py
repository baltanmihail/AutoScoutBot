"""
Генерация профессионального отчёта по стартапу в формате Word (docx).

Используется после «Глубокий анализ»: тот же контент, что в format_deep_analysis_report,
но структурирован в документе с заголовками и абзацами.
"""
from __future__ import annotations

import io
from typing import Any, Dict

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def _add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def build_deep_analysis_docx(analysis: Dict[str, Any]) -> io.BytesIO:
    """
    Строит документ Word из результата глубокого анализа.

    Args:
        analysis: словарь из DeepAnalysisService.analyze_startup_deep_async

    Returns:
        BytesIO с содержимым .docx файла
    """
    doc = Document()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Аналитический отчёт: {analysis.get('startup_name', 'Стартап')}")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()

    internal = analysis.get("internal_analysis", {})

    _add_heading(doc, "1. Технологическая зрелость", level=1)
    tech = internal.get("technology_analysis", {})
    _add_paragraph(doc, f"TRL: {tech.get('trl', 0)}/9")
    _add_paragraph(doc, f"IRL: {tech.get('irl', 0)}/9")
    _add_paragraph(doc, f"MRL: {tech.get('mrl', 0)}/9")
    _add_paragraph(doc, f"CRL: {tech.get('crl', 0)}/9")
    _add_paragraph(doc, f"Средний уровень: {tech.get('average_level', 0):.1f}")
    _add_paragraph(doc, f"Оценка готовности: {tech.get('readiness_assessment', 'н/д')}")
    doc.add_paragraph()

    _add_heading(doc, "2. Финансовый анализ", level=1)
    finance = internal.get("financial_analysis", {})
    avg = finance.get("avg_profit", 0) or 0
    max_p = finance.get("max_profit", 0) or 0
    _add_paragraph(doc, f"Средняя прибыль: {avg / 1_000_000:.2f} млн руб.")
    _add_paragraph(doc, f"Максимальная прибыль: {max_p / 1_000_000:.2f} млн руб.")
    _add_paragraph(doc, f"Тренд: {finance.get('growth_trend', 'н/д')}")
    _add_paragraph(doc, f"Оценка финансового здоровья: {finance.get('financial_health', 'н/д')}")
    doc.add_paragraph()

    recommendations = analysis.get("recommendations", [])
    if recommendations:
        _add_heading(doc, "3. Рекомендации", level=1)
        for i, rec in enumerate(recommendations[:7], 1):
            _add_paragraph(doc, f"{i}. {rec}")
        doc.add_paragraph()

    risks = analysis.get("risk_factors", [])
    if risks:
        _add_heading(doc, "4. Риски", level=1)
        for i, risk in enumerate(risks[:7], 1):
            _add_paragraph(doc, f"{i}. {risk}")
        doc.add_paragraph()

    # Карта рисков (структурированное представление, если появится risk_map)
    risk_map = analysis.get("risk_map", [])
    if isinstance(risk_map, list) and risk_map:
        _add_heading(doc, "4.1. Карта рисков", level=2)
        for item in risk_map[:10]:
            name = item.get("name", "Риск")
            r_type = item.get("type", "")
            prob = item.get("probability", "")
            impact = item.get("impact", "")
            urgency = item.get("urgency", "")
            comment = item.get("comment", "")
            _add_paragraph(
                doc,
                f"{name} — тип: {r_type}, вероятность: {prob}, "
                f"тяжесть: {impact}, срочность: {urgency}. {comment}",
            )
        doc.add_paragraph()

    opportunities = analysis.get("opportunities", [])
    if opportunities:
        _add_heading(doc, "5. Возможности", level=1)
        for i, opp in enumerate(opportunities[:7], 1):
            _add_paragraph(doc, f"{i}. {opp}")
        doc.add_paragraph()

    external = analysis.get("external_analysis", {})
    sources = external.get("sources", [])
    if sources:
        _add_heading(doc, "6. Внешние источники", level=1)
        for src in sources:
            _add_paragraph(doc, f"• {src.get('name', src.get('key', ''))}")
        _add_paragraph(doc, f"Достоверность данных: {external.get('reliability_score', 0):.0%}")
        doc.add_paragraph()

    fin = external.get("financial_data", {})
    if fin:
        _add_heading(doc, "7. Финансы (внешние данные)", level=1)
        for year in sorted(fin.keys(), key=lambda x: int(x) if str(x).isdigit() else 0, reverse=True)[:5]:
            yd = fin[year] if isinstance(fin[year], dict) else {}
            rev = yd.get("revenue", 0) or 0
            profit = yd.get("net_profit", 0) or 0
            rev_s = f"{rev / 1_000_000:.1f} млн руб." if rev else "н/д"
            prof_s = f"{profit / 1_000_000:.1f} млн руб." if profit else "н/д"
            _add_paragraph(doc, f"{year}: выручка {rev_s}, чистая прибыль {prof_s}")
        doc.add_paragraph()

    legal = external.get("legal_status", {})
    if legal:
        _add_heading(doc, "8. Юридический статус (ЕГРЮЛ)", level=1)
        if legal.get("status"):
            _add_paragraph(doc, f"Статус: {legal['status']}")
        if legal.get("registration_date"):
            _add_paragraph(doc, f"Дата регистрации: {legal['registration_date']}")
        doc.add_paragraph()

    news = external.get("news_mentions", [])
    if news:
        _add_heading(doc, "9. Упоминания в СМИ", level=1)
        for mention in news[:5]:
            if mention.get("summary"):
                _add_paragraph(doc, mention["summary"])
            else:
                src = mention.get("source", "").upper()
                title = (mention.get("title", "") or "")[:100]
                _add_paragraph(doc, f"[{src}] {title}")
        doc.add_paragraph()

    smart_articles = analysis.get("smart_articles", [])
    if smart_articles:
        _add_heading(doc, "10. Найденные статьи (AI-поиск)", level=1)
        for art in smart_articles[:5]:
            if art.get("summary"):
                _add_paragraph(doc, art["summary"])
            else:
                _add_paragraph(doc, f"[{art.get('source', '').upper()}] {art.get('title', '')[:100]}")
        doc.add_paragraph()

    doc.add_paragraph()
    footer = doc.add_paragraph("Сформировано AutoScoutBot. Не является инвестиционной рекомендацией.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.italic = True
        run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
