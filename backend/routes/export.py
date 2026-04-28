import io
import pandas as pd
from fastapi.responses import StreamingResponse
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
import json

from backend.database import get_session
from backend.models import Startup, StartupScore, StartupFinancial, ExternalData

router = APIRouter(prefix="/api/export", tags=["export"])

@router.get("/startup/{startup_id}")
async def export_startup_xlsx(startup_id: str, session: AsyncSession = Depends(get_session)):
    stmt = (
        select(Startup, StartupScore)
        .outerjoin(StartupScore, Startup.id == StartupScore.startup_id)
        .where(Startup.id == startup_id)
    )
    row = (await session.execute(stmt)).first()
    if not row:
        raise HTTPException(status_code=404, detail="Startup not found")

    startup, sc = row

    fin_rows = (
        await session.execute(
            select(StartupFinancial)
            .where(StartupFinancial.startup_id == startup.id)
            .order_by(StartupFinancial.year.desc())
        )
    ).scalars().all()

    # Sheet 1: General Info
    general_data = {
        "Показатель": [
            "ID", "Название", "ИНН", "ОГРН", "Год основания", "Статус", "Кластер",
            "Отрасли", "Технологии", "Описание компании",
            "TRL", "IRL", "MRL", "CRL", "Патенты", "Сайт"
        ],
        "Значение": [
            startup.id, startup.name, startup.inn, startup.ogrn, startup.year_founded, startup.status, startup.cluster,
            startup.industries, startup.technologies, startup.company_description,
            startup.trl, startup.irl, startup.mrl, startup.crl, startup.patents, startup.website
        ]
    }
    df_general = pd.DataFrame(general_data)

    # Sheet 2: Financials
    fin_data = []
    for f in fin_rows:
        fin_data.append({"Год": f.year, "Выручка": f.revenue, "Прибыль": f.profit})
    df_fin = pd.DataFrame(fin_data)

    # Sheet 3: ML Scores
    score_data = {"Показатель": [], "Балл (1-10)": []}
    if sc:
        score_data["Показатель"] = [
            "Общий рейтинг (Proxy)", "Зрелость технологий", "Инновационность",
            "Потенциал рынка", "Готовность команды", "Финансовое здоровье",
            "ML Общий балл"
        ]
        score_data["Балл (1-10)"] = [
            sc.score_overall, sc.score_tech_maturity, sc.score_innovation,
            sc.score_market_potential, sc.score_team_readiness, sc.score_financial_health,
            sc.ml_score
        ]
    df_scores = pd.DataFrame(score_data)

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_general.to_excel(writer, sheet_name="Общая информация", index=False)
        df_fin.to_excel(writer, sheet_name="Финансы", index=False)
        df_scores.to_excel(writer, sheet_name="Рейтинг и оценки", index=False)

    output.seek(0)

    filename = f"report_{startup.id}.xlsx"
    from urllib.parse import quote
    filename_safe = quote(f"report_{startup.name}.xlsx".replace(" ", "_").replace('"', ''))
    
    headers = {
        'Content-Disposition': f"attachment; filename*=utf-8''{filename_safe}"
    }

    return StreamingResponse(
        output,
        headers=headers,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@router.get("/startup/{startup_id}/docx")
async def export_startup_docx(startup_id: str, session: AsyncSession = Depends(get_session)):
    stmt = (
        select(Startup, StartupScore)
        .outerjoin(StartupScore, Startup.id == StartupScore.startup_id)
        .where(Startup.id == startup_id)
    )
    row = (await session.execute(stmt)).first()
    if not row:
        raise HTTPException(status_code=404, detail="Startup not found")

    startup, sc = row

    fin_rows = (
        await session.execute(
            select(StartupFinancial)
            .where(StartupFinancial.startup_id == startup.id)
            .order_by(StartupFinancial.year.desc())
        )
    ).scalars().all()

    ext_data = (
        await session.execute(
            select(ExternalData)
            .where(ExternalData.startup_id == startup.id, ExternalData.source == "checko")
            .order_by(ExternalData.fetched_at.desc())
        )
    ).scalars().first()

    news_data = (
        await session.execute(
            select(ExternalData)
            .where(ExternalData.startup_id == startup.id, ExternalData.source == "news")
            .order_by(ExternalData.fetched_at.desc())
        )
    ).scalars().first()

    checko_info = {}
    if ext_data and ext_data.data_json:
        try:
            checko_info = json.loads(ext_data.data_json)
        except:
            pass

    news_info = {}
    if news_data and news_data.data_json:
        try:
            news_info = json.loads(news_data.data_json)
        except:
            pass

    import docx
    from docx.shared import Inches, Pt
    from urllib.parse import quote
    
    doc = docx.Document()
    
    # Title
    title = doc.add_heading(f'Аналитический Отчет: {startup.name}', 0)
    title.alignment = 1 # Center
    
    # Executive Summary
    doc.add_heading('1. Вводное резюме (Executive Summary)', level=1)
    doc.add_paragraph(f'ИНН: {startup.inn or "Н/Д"} | ОГРН: {startup.ogrn or "Н/Д"}')
    doc.add_paragraph(f'Год основания: {startup.year_founded or "Н/Д"}')
    doc.add_paragraph(f'Кластер: {startup.cluster or "Н/Д"} | Статус: {startup.status or "Н/Д"}')
    doc.add_paragraph(f'Описание: {startup.company_description or "Нет описания."}')
    
    # Tech Readiness
    doc.add_heading('2. Технологическая готовность (4 Фазы)', level=1)
    table_tech = doc.add_table(rows=1, cols=2)
    table_tech.style = 'Light Shading Accent 1'
    hdr_cells = table_tech.rows[0].cells
    hdr_cells[0].text = 'Показатель'
    hdr_cells[1].text = 'Уровень (1-9)'
    
    tech_data = [
        ('TRL (Технологическая)', startup.trl),
        ('IRL (Инвестиционная)', startup.irl),
        ('MRL (Производственная)', startup.mrl),
        ('CRL (Коммерческая)', startup.crl)
    ]
    for name, val in tech_data:
        row_cells = table_tech.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(val) if val is not None else "Н/Д"
    
    # Financials
    # 3. Non-Financials
    doc.add_heading('3. Нефинансовые показатели и Команда', level=1)
    if sc:
        team_score = sc.score_team_readiness or 0
        if team_score >= 8:
            team_eval = "Высокая готовность. Опытная команда с релевантным бэкграундом."
        elif team_score >= 5:
            team_eval = "Средняя готовность. Требуется усиление отдельных компетенций."
        else:
            team_eval = "Низкая готовность или недостаточно данных о команде."
        
        inn_score = sc.score_innovation or 0
        if inn_score >= 8:
            inn_eval = "Высокая степень инновационности, потенциально прорывные технологии."
        else:
            inn_eval = "Умеренная или низкая степень инновационности."
            
        doc.add_paragraph(f'Оценка Команды: {team_score}/10 — {team_eval}')
        doc.add_paragraph(f'Инновационность: {inn_score}/10 — {inn_eval}')
        doc.add_paragraph(f'Рыночный потенциал: {sc.score_market_potential or "Н/Д"}/10')
    else:
        doc.add_paragraph("Данные ML-оценки отсутствуют.")

    # 4. News & Media
    doc.add_heading('4. Инфополе и Новости', level=1)
    if news_info and "articles" in news_info and news_info["articles"]:
        doc.add_paragraph('Автоматический анализ упоминаний в СМИ (Google News RSS):')
        for i, article in enumerate(news_info["articles"][:5]):
            pub_date = article.get("published_at", "")
            title = article.get("title", "")
            source = article.get("source", "")
            doc.add_paragraph(f"{i+1}. [{pub_date}] {title} ({source})")
        doc.add_paragraph('• Тональность инфополя: требует ручной верификации.')
    else:
        doc.add_paragraph('Автоматический анализ упоминаний в СМИ:')
        doc.add_paragraph('• На данный момент значимых новостных публикаций в федеральных СМИ не найдено.')
        doc.add_paragraph('• Тональность инфополя: нейтральная.')
        doc.add_paragraph('• PR-активность: требуется дополнительное исследование.')

    # 5. Legal Due Diligence
    doc.add_heading('5. Юридический Due Diligence', level=1)
    if checko_info:
        legal = checko_info.get("legal_cases", {})
        enf = checko_info.get("enforcements", {})
        contracts = checko_info.get("contracts", {})
        
        if legal:
            doc.add_paragraph(f'Судебные дела (Арбитраж): Найдено записей. (Подробнее в полной выписке Checko).')
        else:
            doc.add_paragraph('Судебные дела (Арбитраж): Нет активных дел или данных.')
            
        if enf:
            doc.add_paragraph(f'Исполнительные производства (ФССП): Найдено записей. (Подробнее в полной выписке Checko).')
        else:
            doc.add_paragraph('Исполнительные производства (ФССП): Задолженностей не найдено.')
            
        if contracts:
            doc.add_paragraph('Госконтракты (ФЗ-44, ФЗ-223): Компания участвовала в госзакупках.')
        else:
            doc.add_paragraph('Госконтракты: Данных об участии не найдено.')
    else:
        doc.add_paragraph("Расширенные юридические данные (Checko) не загружены для данной компании.")

    # 6. Financials
    doc.add_heading('6. Финансовый аудит и метрики', level=1)
    if not fin_rows:
        doc.add_paragraph("Финансовые данные отсутствуют в базе.")
    else:
        table_fin = doc.add_table(rows=1, cols=3)
        table_fin.style = 'Light Shading Accent 1'
        hdr_cells = table_fin.rows[0].cells
        hdr_cells[0].text = 'Год'
        hdr_cells[1].text = 'Выручка (тыс. руб)'
        hdr_cells[2].text = 'Прибыль (тыс. руб)'
        for f in fin_rows:
            row_cells = table_fin.add_row().cells
            row_cells[0].text = str(f.year)
            row_cells[1].text = f"{f.revenue:,.0f}" if f.revenue is not None else "Н/Д"
            row_cells[2].text = f"{f.profit:,.0f}" if f.profit is not None else "Н/Д"
            
    # 7. Risk Matrix
    doc.add_heading('7. Матрица рисков и Оценка', level=1)
    doc.add_paragraph('Оценка рисков проведена на основе доступных метрик и ML-моделирования.')
    if sc:
        doc.add_paragraph(f'Общий ML Рейтинг: {sc.score_overall or "Н/Д"} / 10')
        doc.add_paragraph(f'Риск банкротства (Z-Score): Анализ в дашборде.')
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    filename_safe = quote(f"report_{startup.name}.docx".replace(" ", "_").replace('"', ''))
    headers = {
        'Content-Disposition': f"attachment; filename*=utf-8''{filename_safe}"
    }

    return StreamingResponse(
        output,
        headers=headers,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )